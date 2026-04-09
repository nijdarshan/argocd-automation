"""
Post-process the docx to ensure ALL tables have visible borders.
Pandoc sometimes ignores template table styles — this forces borders on every table.
"""
from docx import Document
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

BORDER_COLOR = "666666"
BORDER_SZ = "6"

def add_borders_to_table(table):
    """Add visible borders to a table element."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)

    # Remove existing borders
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)

    # Add new borders
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{BORDER_SZ}" w:space="0" w:color="{BORDER_COLOR}"/>'
        f'  <w:left w:val="single" w:sz="{BORDER_SZ}" w:space="0" w:color="{BORDER_COLOR}"/>'
        f'  <w:bottom w:val="single" w:sz="{BORDER_SZ}" w:space="0" w:color="{BORDER_COLOR}"/>'
        f'  <w:right w:val="single" w:sz="{BORDER_SZ}" w:space="0" w:color="{BORDER_COLOR}"/>'
        f'  <w:insideH w:val="single" w:sz="{BORDER_SZ}" w:space="0" w:color="{BORDER_COLOR}"/>'
        f'  <w:insideV w:val="single" w:sz="{BORDER_SZ}" w:space="0" w:color="{BORDER_COLOR}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

    # Bold the first row (header)
    first_row = table.rows[0] if table.rows else None
    if first_row:
        for cell in first_row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True

docx_path = '/Users/darsh/vmo2/argocd-automation/VMO2_Application_Onboarding_Technical_Design.docx'
doc = Document(docx_path)

count = 0
for table in doc.tables:
    add_borders_to_table(table)
    count += 1

doc.save(docx_path)
print(f"Fixed borders on {count} tables")
